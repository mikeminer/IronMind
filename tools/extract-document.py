import json
import os
import sys


def emit(value):
    sys.stdout.write(json.dumps(value, ensure_ascii=False))


def text_sections(text, source="text"):
    sections = []
    para = 0
    for block in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n\n"):
        clean = " ".join(block.split())
        if not clean:
            continue
        para += 1
        sections.append({"source": source, "page": None, "paragraph": para, "text": clean})
    if not sections and str(text or "").strip():
        sections.append({"source": source, "page": None, "paragraph": 1, "text": " ".join(str(text).split())})
    return sections


def extract_txt(file_path):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as handle:
                return text_sections(handle.read(), "text")
        except UnicodeDecodeError:
            continue
    with open(file_path, "rb") as handle:
        return text_sections(handle.read().decode("latin-1", errors="replace"), "text")


def extract_pdf(file_path):
    warnings = []
    sections = []
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                paragraphs = text_sections(text, "pdf")
                for item in paragraphs:
                    item["page"] = page_index
                sections.extend(paragraphs)
        if sections:
            return sections, warnings
    except Exception as exc:
        warnings.append(f"pdfplumber failed: {exc}")

    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            paragraphs = text_sections(text, "pdf")
            for item in paragraphs:
                item["page"] = page_index
            sections.extend(paragraphs)
    except Exception as exc:
        warnings.append(f"pypdf failed: {exc}")

    return sections, warnings


def extract_docx(file_path):
    from docx import Document

    doc = Document(file_path)
    sections = []
    paragraph = 0
    for item in doc.paragraphs:
        clean = " ".join((item.text or "").split())
        if clean:
            paragraph += 1
            sections.append({"source": "docx", "page": None, "paragraph": paragraph, "text": clean})
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [" ".join((cell.text or "").split()) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                rows.append(line)
        if rows:
            paragraph += 1
            sections.append({
                "source": "docx",
                "page": None,
                "paragraph": paragraph,
                "text": f"Tabella {table_index}: " + " / ".join(rows)
            })
    return sections


def main():
    if len(sys.argv) < 2:
        emit({"ok": False, "error": "missing file path"})
        return 2

    file_path = sys.argv[1]
    ext = os.path.splitext(file_path)[1].lower()
    warnings = []
    try:
        if ext == ".pdf":
            sections, warnings = extract_pdf(file_path)
        elif ext == ".docx":
            sections = extract_docx(file_path)
        elif ext in (".txt", ".md", ".markdown", ".csv"):
            sections = extract_txt(file_path)
        else:
            emit({"ok": False, "error": f"unsupported file extension: {ext}"})
            return 2

        emit({
            "ok": True,
            "sections": sections,
            "warnings": warnings,
            "characters": sum(len(item.get("text", "")) for item in sections)
        })
        return 0
    except Exception as exc:
        emit({"ok": False, "error": str(exc), "warnings": warnings})
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
